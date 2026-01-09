import os
from dotenv import load_dotenv
from google import genai

# DOCX
from docx import Document
from docx.shared import Pt

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch


# ======================================================
# LOAD ENV VARIABLES (SAFE FOR IMPORT)
# ======================================================
load_dotenv()

API_KEY = os.getenv("GOOGLE_API_KEY")
client = genai.Client()

MODEL_NAME = "gemini-2.5-flash"


# ======================================================
# PROMPT BUILDER
# ======================================================
def build_general_legal_prompt(user_input: str) -> str:
    """
    Builds a prompt for generating legal documents
    """
    return f"""
ROLE:
You are an expert Indian legal draftsman specializing in drafting GENERAL, NEUTRAL, and PROCEDURAL complaints for Indian authorities such as Consumer Commissions, Civil Courts, Police Authorities, Cyber Crime Cells, Family Courts, and other statutory forums.

Your task is to generate a complaint in a FIXED FOUR-PAGE FORMAT.
The output must be suitable for filing and must NOT presume guilt, offence, defect, or liability unless explicitly stated by the user.

────────────────────────────────
MANDATORY RULES (STRICT)
────────────────────────────────
1. Use ONLY facts explicitly stated by the user.
2. DO NOT assume, infer, expand, or fabricate:
   - Facts, incidents, dates, amounts
   - Legal sections, offences, or conclusions
3. If information is missing:
   - Keep the content general and procedural
   - Do NOT insert placeholders, brackets, or commentary
4. Maintain formal, neutral, court-acceptable language.
5. DO NOT accuse, define scams, or conclude wrongdoing.
6. Output ONLY the complaint document.
7. DO NOT use stars, separators, markdown, or explanations.
8. ALWAYS generate EXACTLY FOUR PAGES.
9. Preserve page meaning and order strictly.

────────────────────────────────
PAGE DEFINITIONS (MUST FOLLOW)
────────────────────────────────

PAGE 1:
- Authority heading
- Details of Complainant
- Details of Opponent
- NO narration
- NO facts
- NO prayers

PAGE 2:
- Instructional and conditional framework
- Use phrases like “If the complaint is regarding…”
- Explain WHAT details may be stated, not WHAT happened
- Adapt vocabulary based on complaint type (consumer, civil, cyber, domestic)
- NO assertions

PAGE 3:
- Statement regarding other proceedings
- Limitation or delay condonation (either-or)
- Conditional list of possible prayers
- Fees details
- NO mandatory reliefs

PAGE 4:
- Authorized representative details (if any)
- Declaration of truth
- Place and signature
- Enclosures list

────────────────────────────────
FOUR PAGE TEMPLATE
────────────────────────────────

FORMAT OF (CONSUMER/CIVIL/DOMESTIC/CYBER/ETC)COMPLAINT

BEFORE THE HON’BLE ____________________________________________________________

1. Details of the Complainant

A. Full Name:
B. Full Address:                                                                                                
C. Mobile No.:
D. E-Mail:

If there is more than one complainant, then provide information of all complainants for Sr. No. (A to D) above.

V/s.

2. Details of the Opponent

A) Full Name:
B) Full Address (With PIN Code No.):                                                                          
C) Phone / Mobile No.:
D) E-Mail Address:

If there is more than one opponent, then provide information of all opponents for Sr. No. (A to D) above.



I/We, the complainant(s), humbly request to the Hon’ble Authority that:

(1) Submit all the details of the complaint i.e. how, when, and where the cause of action arose.
The events should be mentioned in chronological order, based strictly on the complainant’s information.
A. If the complaint relates to goods or services, provide details of the nature of goods or services involved, if applicable.
B. If the complaint relates to any transaction, activity, conduct, or relationship, provide relevant details thereof, if applicable.
C. Provide details of any amount involved, if any, along with supporting documents, if available.
D. Provide details of documents relied upon, if any.

(2) A conditional instruction that applies only if the complaint relates to unfair trade practice.
Include sub-points:
A.An instructional line asking the complainant to state the type or nature of unfair trade practice, if applicable.
B.An instructional line asking the complainant to state details of any damage, loss, or inconvenience suffered, if any, along with supporting documents, if available.
Ensure the language remains conditional and procedural.
Do not assert that any unfair trade practice has occurred unless explicitly stated by the user.

(3) An instructional and conditional clause requesting details of efforts made by the complainant to resolve the matter, if any.
Include sub-points:
A.Verbal communication or attempts, if any.
B.Written correspondence, notice, or communication, if any.
C.For e-commerce related matters, token number, reference number, or complaint ID provided by the platform, if any.
D.Result or outcome of such efforts, if any, along with documentary evidence, if available.
Ensure the language remains conditional and does not assume that any efforts were made unless stated by the complainant.

(4) An instructional and conditional clause stating that the District or State Commission has jurisdiction, if applicable.
Include sub-points:
A. A conditional instruction to state jurisdiction based on cause of action, residence of the complainant, or place of business or residence of the opponent, if applicable.
B. A conditional instruction to state pecuniary jurisdiction based on the claim amount, distinguishing between District and State Commission limits.
Include the phrase “Strike out whichever is not applicable.”
Ensure that jurisdiction is not asserted as a fact unless explicitly supported by information provided by the complainant.

(5) An instructional declaration requiring the complainant to state that no complaint regarding the present matter has been filed before any other Court, Tribunal, or Commission.
If any complaint has been filed earlier, instruct the complainant to provide details of such complaint, the outcome thereof, and documentary evidence, if any.
Ensure the language remains declaratory and conditional, and does not assume the existence of prior proceedings.

(6) Any other details which the complainant wishes to submit.

(7) A conditional statement that the complaint is filed within the limitation period, where applicable.
                                                       OR
If the complaint is time barred, include a statement that a delay condonation application for delay of __________ days is attached.
Ensure the language remains procedural and does not assume whether the complaint is within limitation or time barred.

(8) An instructional clause stating that if the complainant wishes to seek any relief, the same may be mentioned.
Include sub-points:
(A) Refund of price or charges paid, if applicable.
(B) Replacement of goods with new goods, if applicable.
(C) Removal of defects in goods, if applicable.
(D) Compensation for deficiency in service or negligence, if applicable.
(E) Compensation for mental torture and cost of the complaint, if applicable.
(F) Payment of unpaid or less paid insurance claim amount with interest, if applicable.
(G) Any other relief deemed fit and proper.
Include a concluding note stating that if compensation is claimed, the calculation of compensation sought should be provided.
Ensure the language remains conditional and does not assume entitlement to any relief.

(9) A heading stating “Details of fees paid at the time of filing the complaint”.
Below the heading, include separate lines for:
- Claim Amount (Rs.):                                    Fees (Rs.):
- Demand Draft Number:                                   Date:
- RTGS / NEFT:
Do not assume that any fee has been paid.
Leave the fields blank unless information is explicitly provided by the complainant.


(10) If the complaint is filed through an authorized representative, advocate, or association, provide details.
Name:
Full Address:
Mobile No.:
E-Mail:

(11) Declaration
I, ____________, hereby declare that the above mentioned information is true and correct to the best of my knowledge and belief.



Place:   
                                                                                                               

Signature of Complainant 


Enclosures:
1. Proof of amount deposited, if any
2. List of documentary evidences, if any
3. Vakalatnama, if Advocate is engaged

────────────────────────────────
USER QUERY
────────────────────────────────
{user_input}

────────────────────────────────
FINAL INSTRUCTION
────────────────────────────────
Generate the complaint strictly in the above FOUR-PAGE format, preserving neutrality, structure, and procedural tone.

"""


# ======================================================
# GEMINI GENERATION (SYNCHRONOUS)
# ======================================================
def generate_legal_text(user_input: str) -> str:
    """
    Generate legal document text using Gemini AI
    
    Args:
        user_input (str): User's request with facts and details
        
    Returns:
        str: Generated legal document text
    """
    try:
        print(f"Calling Gemini API with model: {MODEL_NAME}")
        
        # model = genai.GenerativeModel(MODEL_NAME)
        prompt = build_general_legal_prompt(user_input)
        # Generate content
        response = client.models.generate_content(
            model = "gemini-2.5-flash",
            contents = prompt,
            config={
                "temperature": 0.3,
                "top_p": 0.8,
                "top_k": 40,
                "max_output_tokens": 8192,
            }
        )

        # Extract text
        legal_text = response.text.strip()
        
        print(f"✓ Successfully generated legal document ({len(legal_text)} characters)")
        
        return legal_text
        
    except Exception as e:
        print(f"✗ Error generating legal text: {str(e)}")
        raise RuntimeError(f"Failed to generate legal document: {str(e)}")


# ======================================================
# DOCX CREATOR
# ======================================================
def save_to_docx(text: str, filename: str) -> str:
    """
    Save legal text to a DOCX file
    
    Args:
        text (str): Legal document text
        filename (str): Output filename (with .docx extension)
        
    Returns:
        str: Path to saved file
    """
    try:
        doc = Document()
        
        # Set default style
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        
        # Add content with proper spacing
        for line in text.split("\n"):
            if line.strip():  # Skip empty lines
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)
            else:
                doc.add_paragraph()  # Add blank line
        
        # Save document
        doc.save(filename)
        print(f"✓ DOCX saved: {filename}")
        
        return filename
        
    except Exception as e:
        print(f"✗ Error saving DOCX: {str(e)}")
        raise RuntimeError(f"Failed to save DOCX file: {str(e)}")


# ======================================================
# PDF CREATOR
# ======================================================
def save_to_pdf(text: str, filename: str) -> str:
    """
    Save legal text to a PDF file
    
    Args:
        text (str): Legal document text
        filename (str): Output filename (with .pdf extension)
        
    Returns:
        str: Path to saved file
    """
    try:
        # Create PDF document
        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            rightMargin=1 * inch,
            leftMargin=1 * inch,
            topMargin=1 * inch,
            bottomMargin=1 * inch,
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        story = []
        
        # Process text line by line
        for line in text.split("\n"):
            if line.strip():  # Skip empty lines
                # Escape special XML characters
                safe_line = (
                    line.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                        .replace('"', "&quot;")
                        .replace("'", "&apos;")
                )
                story.append(Paragraph(safe_line, styles["Normal"]))
            else:
                story.append(Paragraph("<br/>", styles["Normal"]))  # Blank line
        
        
        doc.build(story)
        print(f"✓ PDF saved: {filename}")
        
        return filename
        
    except Exception as e:
        print(f"✗ Error saving PDF: {str(e)}")
        raise RuntimeError(f"Failed to save PDF file: {str(e)}")



def generate_and_save(user_input: str, output_dir: str = "outputs") -> dict:
    """
    Generate legal document and save to both DOCX and PDF
    
    Args:
        user_input (str): User's request with facts
        output_dir (str): Directory to save files
        
    Returns:
        dict: Contains text and file paths
    """
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate legal text
        legal_text = generate_legal_text(user_input)
        
        # Create filenames
        base_name = "legal_document"
        docx_path = os.path.join(output_dir, f"{base_name}.docx")
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        
        # Save to both formats
        save_to_docx(legal_text, docx_path)
        save_to_pdf(legal_text, pdf_path)
        
        return {
            "text": legal_text,
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "success": True
        }
        
    except Exception as e:
        print(f"✗ Error in generate_and_save: {str(e)}")
        return {
            "text": None,
            "docx_path": None,
            "pdf_path": None,
            "success": False,
            "error": str(e)
        }